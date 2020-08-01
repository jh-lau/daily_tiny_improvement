"""
  @Author       : Liujianhan
  @Date         : 20/7/20 21:56
  @FileName     : scanner.py
  @ProjectName  : DailyTinyImprovement
  @Description  : Placeholder
 """
import codecs
import os
import re
import shutil
import subprocess
import time
import logging
from concurrent import futures
from sys import platform
from typing import List
import pandas as pd
from io import StringIO
from io import open
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfinterp import PDFResourceManager, process_pdf
from PIL import Image
import pytesseract
from docx import Document
import configparser
from collections import defaultdict

logging.propagate = False
logging.getLogger().setLevel(logging.ERROR)

# config_list = yaml.safe_load(codecs.open('snippets/scanner_config.yml'))
# try:
#     target_file_ext = config_list['target_file_type']
#     key_word_list = config_list['target_word']
#     target_root_dir = config_list['target_root_dir']
#     output_dir = config_list['output_dir']
#     buffer_size = int(config_list['buffer_size'])
try:
    config_reader = configparser.ConfigParser()
    config_reader.read('config.ini')
    key_word_list = config_reader.get('INPUT', 'key').split(',')
    buffer_size = int(config_reader.get('PROCESS', 'buff'))
    target_root_dir = config_reader.get('INPUT', 'path').split(',')
    target_file_ext = config_reader.get('INPUT', 'file_type').split(',')
    # target_file_ext = tuple([f'.{ext}' for ext in target_file_ext])
    output_dir = 'output'
except KeyError as e:
    print(f'配置文件信息 {e} 未找到，请检查。')
rsrcmgr = PDFResourceManager()

file_type_map = {
    'txt': ('.tmp', '.log', '.txt', '.md'),
    'excel': ('.xls', '.xlsx'),
    'pdf': ('.pdf',),
    'doc': ('.doc',),
    'docx': ('.docx',),
    'pics': ('.jpg', '.png', '.bmp', '.jpeg')
}


def file_counter(file_path: str):
    inner_counter = 0
    # target_file = []
    file_dic = defaultdict(set)
    for root, _, files in os.walk(file_path):
        for file in files:
            file = os.path.join(root, file)
            if file.endswith(file_type_map['txt']):
                file_dic['txt'].add(file)
            elif file.endswith(file_type_map['excel']):
                file_dic['excel'].add(file)
            elif file.endswith(file_type_map['pdf']):
                file_dic['pdf'].add(file)
            elif file.endswith(file_type_map['pics']):
                file_dic['pics'].add(file)
            elif file.endswith(file_type_map['doc']):
                file_dic['doc'].add(file)
            elif file.endswith(file_type_map['docx']):
                file_dic['docx'].add(file)
    inner_counter += sum(map(len, file_dic.values()))

    # return inner_counter, target_file
    return inner_counter, file_dic


def file_checker(file_path: str):
    try:
        result = ''
        if file_path.endswith(file_type_map['txt']):
            result = txt_checker(file_path)
        elif file_path.endswith(file_type_map['excel']):
            result = excel_checker(file_path)
        elif file_path.endswith(file_type_map['pdf']):
            result = pdf_checker(file_path)
        elif file_path.endswith(file_type_map['pics']):
            result = pics_checker(file_path)
        elif file_path.endswith(file_type_map['doc']):
            result = doc_checker(file_path)
        elif file_path.endswith(file_type_map['docx']):
            result = docx_checker(file_path)

        return result
    except Exception as e:
        pass


def txt_checker(file_path):
    with open(file_path, 'r', encoding="utf-8") as f:
        curr_ = f.read(buffer_size)
        next_ = f.read(buffer_size)
        while True:
            if not curr_ and not next_:
                break
            if search_keyword("%s%s" % (curr_, next_)):
                return file_path
            curr_ = next_
            next_ = f.read(buffer_size)


def excel_checker(file_path):
    excel = pd.ExcelFile(file_path)
    for a_sheet_name in excel.sheet_names:
        d = pd.read_excel(file_path, sheet_name=a_sheet_name, dtype=str)
        ri, ci = d.shape
        if ri == 0 and ci == 0:
            continue
        for title in d.columns:
            # 处理
            if title and search_keyword(str(title)):
                return file_path
        for row in d.values:
            for col in row:
                if col and search_keyword(str(col)):
                    return file_path


def pdf_checker(file_path):
    with open(file_path, "rb") as pdf:
        retstr = StringIO()
        device = TextConverter(rsrcmgr, retstr, laparams=LAParams())
        process_pdf(rsrcmgr, device, pdf)
        device.close()
        content = retstr.getvalue()
        retstr.close()
        if search_keyword(content.replace("\n", "")):
            return file_path


def pics_checker(file_path):
    if search_keyword(
            ''.join(''.join(pytesseract.image_to_string(Image.open(file_path), lang='chi_sim').split()).strip())):
        return file_path


def docx_checker(file_path):
    document = Document(file_path)
    # 文档内容
    for paragraph in document.paragraphs:
        if search_keyword(paragraph.text):
            return file_path
    # 文档表格
    tables = document.tables
    for table in tables[:]:
        for i, row in enumerate(table.rows[:]):  # 读每行
            for cell in row.cells[:]:  # 读一行中的所有单元格
                if search_keyword(cell.text):
                    return file_path


def doc_checker(file_path):
    if platform.startswith("win"):
        pass
    else:
        if search_keyword(subprocess.check_output(['antiword', file_path]).decode("utf-8")):
            return file_path


def gen_path(target_root_path: str):
    count = 0
    for file in os.listdir(target_root_path):
        print(f'\r处理第 {count} 个路径。', end='')
        yield f'/{file}'
        count += 1


def gen_file(file_list):
    count = 0
    for file in file_list:
        print(f'\r搜集第 {count} 个文件.', end='')
        yield file
        count += 1


def copy_file(file_path: str, output_dir: str = output_dir):
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    base_name = os.path.basename(file_path)
    shutil.copy(file_path, f'{output_dir}/{base_name}')


def search_keyword(stream_content: str,
                   key_word_list: List = key_word_list):
    # 匹配关键词中间带空格的情况
    key_word_list = ['\s*'.join(s) for s in [list(s) for s in key_word_list]]
    pattern = '|'.join(key_word_list)
    result = re.findall(f'{pattern}', stream_content)
    return True if result else False


if __name__ == '__main__':
    path_list = gen_path('/')
    start = time.time()
    counter = 0
    file_type_list = ['txt', 'pdf', 'excel', 'doc', 'docx', 'pics']
    total_files_dic = {s: set() for s in file_type_list}
    print(f'开始从 {target_root_dir} 检测：目标文件类型为{target_file_ext}')

    with futures.ProcessPoolExecutor() as pool:
        for number, target_file_dic in pool.map(file_counter, path_list):
            counter += number
            for key in total_files_dic:
                total_files_dic[key].update(target_file_dic.get(key, set()))
    print(f'\n统计文件耗时 : {time.time() - start:.3f}秒， 共有待处理文件 {counter} 个。')
    for key, value in total_files_dic.items():
        print(f'共有 {key} 文件： {len(value)} 个。')

    start = time.time()
    hit_file = set()
    for key, value in total_files_dic.items():
        count = 0
        print(f'\n开始处理 {key} 类型文件...')
        file_gen = gen_file(value)
        inside_timer = time.time()
        with futures.ProcessPoolExecutor() as pool:
            results = pool.map(file_checker, file_gen)
            print()
            for result in results:
                if result:
                    hit_file.add(result)
                count += 1
                elapse_time = (time.time() - inside_timer) / 60
                print(f'\r处理第 {count} 个文件，累计耗时： {elapse_time:<6.1f} 分钟', end='')
    print(f'\n处理所有文件耗时 : {(time.time() - start) / 60:.1f}分钟。', end='')

    if hit_file:
        with futures.ProcessPoolExecutor() as pool:
            pool.map(copy_file, hit_file)
        print(f'检测到涉密文件 {len(hit_file)} 个，已复制到当前 {output_dir} 输出文件夹。')
    else:
        print(f'通过检测，未发现包含敏感信息文件。')
